from __future__ import annotations

import re
from pathlib import Path
from typing import Literal

from docx import Document

from .db import DishKind

# Heuristic: Russian menu keywords
_GARNISH_RE = re.compile(
    r"гарнир|каш[ае]|гречк|рис\b|макарон|пюре|картоф|овощ|капуст|фасол|горох|перлов",
    re.IGNORECASE,
)
_MAIN_RE = re.compile(
    r"котлет|отбивн|тефтел|биток|мяс|рыб|курин|индейк|печен|печён|филе|"
    r"гуляш|подлив|жаркое|бефстр|стейк|свин|говя|баран|фрикадел|рулет",
    re.IGNORECASE,
)


def classify_dish(name: str) -> DishKind:
    n = name.strip()
    g = bool(_GARNISH_RE.search(n))
    m = bool(_MAIN_RE.search(n))
    if g and not m:
        return "garnish"
    if m and not g:
        return "main"
    if g and m:
        return "main"
    return "other"


_CALORIES_IN_NAME_RE = re.compile(
    r"""
    (?:
        \s*[\(\[]?\s*
        \d{1,4}\s*(?:ккал(?:орий)?|kcal)\b
        \s*[\)\]]?
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


def strip_calories_from_dish_name(name: str) -> str:
    n = name.strip()
    if not n:
        return n
    prev = None
    while prev != n:
        prev = n
        n = _CALORIES_IN_NAME_RE.sub(" ", n)
        n = re.sub(r"\s+", " ", n).strip()
    return n


def _is_nutrition_numeric_token(t: str) -> bool:
    t = t.strip()
    if not t:
        return False
    if re.fullmatch(r"\d+шт", t, re.IGNORECASE):
        return True
    if re.fullmatch(r"\d+[,.]\d+", t):
        return True
    if re.fullmatch(r"\d+", t):
        return True
    if re.fullmatch(r"\d+/\d+", t):
        return True
    if re.fullmatch(r"\d+,", t):
        return True
    return False


def strip_nutrition_table_tail(name: str) -> str:
    parts = name.split()
    if len(parts) < 4:
        return name
    i = len(parts)
    while i > 0 and _is_nutrition_numeric_token(parts[i - 1]):
        i -= 1
    removed = len(parts) - i
    if removed >= 4:
        return " ".join(parts[:i]).strip()
    return name


def sanitize_dish_name(name: str) -> str:
    n = name.strip()
    if not n:
        return n
    n = strip_nutrition_table_tail(n)
    n = strip_calories_from_dish_name(n)
    n = re.sub(r"\s+", " ", n).strip()
    return n


_PRICE_RE = re.compile(
    r"""
    (?P<name>.+?)
    [\s\u00a0]*                           # spaces
    (?P<price>\d+(?:[.,]\d{1,2})?)       # price
    \s*(?:руб|р\.?|₽)?\s*$               # optional rub
    """,
    re.VERBOSE | re.IGNORECASE,
)


def _parse_price_token(token: str) -> float | None:
    s = (token or "").strip()
    if not s:
        return None
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+", "", s)
    s = s.strip(".,;:()[]{}")
    s = s.rstrip("рубРУБ.р₽")
    m_dash = re.fullmatch(r"(\d+)-(\d{1,2})", s)
    if m_dash:
        rub = m_dash.group(1)
        kop = m_dash.group(2).ljust(2, "0")
        s = f"{rub}.{kop}"
    else:
        s = s.replace(",", ".")
    if not re.fullmatch(r"\d+(?:\.\d{1,2})?", s):
        return None
    try:
        price = float(s)
    except ValueError:
        return None
    if price <= 0 or price > 100000:
        return None
    return price


_CATEGORY_PREFIXES_MULTI = frozenset(
    {
        "выпечка",
        "гарниры",
        "напитки",
        "второе",
        "вторые",
        "десерты",
        "закуски",
        "блюда",
    }
)
_CATEGORY_PREFIXES_SINGLE = frozenset({"выпечка", "гарниры", "напитки"})

_DASH_PRICE_IN_LINE_RE = re.compile(r"(?<!\S)(\d+-\d{1,2})(?!\S)")


def _maybe_strip_category_prefix(name: str, *, multi_item_line: bool) -> str:
    parts = name.split()
    if len(parts) < 2:
        return name
    first = parts[0].casefold()
    allowed = _CATEGORY_PREFIXES_MULTI if multi_item_line else _CATEGORY_PREFIXES_SINGLE
    if first in allowed:
        return " ".join(parts[1:]).strip()
    return name


def split_multi_price_line(line: str) -> list[tuple[str, float]] | None:
    s = line.strip()
    if not s:
        return None
    s = re.sub(r"\s*[—–]\s*", " ", s)
    s = re.sub(r"\s*₽\s*$", "", s)
    s = re.sub(r"\s*руб\.?\s*$", "", s, flags=re.IGNORECASE)
    matches = list(_DASH_PRICE_IN_LINE_RE.finditer(s))
    if not matches:
        return None
    out: list[tuple[str, float]] = []
    prev = 0
    for m in matches:
        tok = m.group(1)
        price = _parse_price_token(tok)
        if price is None:
            continue
        chunk = s[prev : m.start()].strip()
        prev = m.end()
        if chunk:
            out.append((chunk, price))
    return out if out else None


def _split_top_level_by_commas(s: str) -> list[str]:
    parts: list[str] = []
    buf: list[str] = []
    depth = 0
    pairs = {"(": ")", "[": "]", "{": "}"}
    closing = {v: k for k, v in pairs.items()}
    for ch in s:
        if ch in pairs:
            depth += 1
        elif ch in closing:
            depth = max(0, depth - 1)
        if depth == 0 and ch in [",", ";"]:
            part = "".join(buf).strip()
            if part:
                parts.append(part)
            buf = []
        else:
            buf.append(ch)
    tail = "".join(buf).strip()
    if tail:
        parts.append(tail)
    return parts


# Объём/вес «0,5» или «10,5» — не разделитель списка блюд при split по запятым.
_VOL_DECIMAL_COMMA_RE = re.compile(r"(?<![\d,])(\d+,\d{1,2})(?!\d)")


def _protect_volume_decimal_commas(s: str) -> tuple[str, dict[str, str]]:
    """Временно убирает запятую в десятичных дробях (0,5 л), чтобы _split_top_level_by_commas не резал имя."""
    tokens: dict[str, str] = {}
    n = 0

    def repl(m: re.Match[str]) -> str:
        nonlocal n
        key = f"\x00dc{n}\x00"
        tokens[key] = m.group(1)
        n += 1
        return key

    return _VOL_DECIMAL_COMMA_RE.sub(repl, s), tokens


def _restore_volume_decimal_commas(s: str, tokens: dict[str, str]) -> str:
    for k, v in tokens.items():
        s = s.replace(k, v)
    return s


_EXPLICIT_RUB_TOKEN_RE = re.compile(
    r"\d+(?:[.,]\d{1,2})?\s*(?:руб|р\.?|р)\b",
    re.IGNORECASE,
)


def _split_name_and_trailing_number_token(name: str) -> tuple[str, int | None]:
    m = re.search(r"^(.*\D)\s+(\d{1,3})\s*$", name)
    if not m:
        return name.strip(), None
    base = m.group(1).strip()
    if not base:
        return name.strip(), None
    return base, int(m.group(2))


def _has_explicit_price_marker(text: str) -> bool:
    return bool(_EXPLICIT_RUB_TOKEN_RE.search(text) or re.search(r"\d+-\d{1,2}", text))


def split_multi_comma_price_line(line: str) -> list[tuple[str, float]] | None:
    if not line or not _EXPLICIT_RUB_TOKEN_RE.search(line):
        return None

    s = line.strip()
    s = s.strip(" \t\r\n")

    s_masked, dc_toks = _protect_volume_decimal_commas(s)
    raw_parts = _split_top_level_by_commas(s_masked)
    out: list[tuple[str, float]] = []

    for part in raw_parts:
        part = _restore_volume_decimal_commas(part.strip(), dc_toks)
        part = part.strip().strip(".,;:")
        if not part:
            continue
        part = re.sub(r"^\s*\d+\s*[.)]\s*", "", part)

        if not _EXPLICIT_RUB_TOKEN_RE.search(part):
            continue

        parsed = _parse_line(part)
        if parsed is None:
            continue
        name, price = parsed
        if name and price is not None:
            out.append((name, price))

    return out if len(out) >= 2 else None


def split_multi_comma_mixed_price_line(line: str) -> list[tuple[str, float]] | None:
    """
    Список через запятую, где у части позиций цена указана явно, а у части — "голым" числом:
      "Пицца 100р, Чебурек 70р, Сосиска в тесте 60, ...";
      "Хот-дог 80, Сочень 50".
    """
    if not line or "," not in line:
        return None

    s = line.strip().strip(" \t\r\n")
    s_masked, dc_toks = _protect_volume_decimal_commas(s)
    raw_parts = _split_top_level_by_commas(s_masked)
    if len(raw_parts) < 2:
        return None

    out: list[tuple[str, float]] = []
    explicit_count = 0
    trailing_count = 0
    for part in raw_parts:
        p = _restore_volume_decimal_commas(part.strip(), dc_toks)
        p = p.strip().strip(".,;:")
        p = re.sub(r"^\s*\d+\s*[.)]\s*", "", p)
        if not p:
            continue

        parsed = _parse_line(p)
        if parsed is not None:
            name, price = parsed
            out.append((name, price))
            if _has_explicit_price_marker(p):
                explicit_count += 1
            else:
                trailing_count += 1
            continue

        base_name, trailing_num = _split_name_and_trailing_number_token(p)
        if trailing_num is not None and 5 <= trailing_num <= 1000:
            out.append((base_name, float(trailing_num)))
            trailing_count += 1

    # Защита от ложных срабатываний:
    # - либо есть хотя бы одна явная цена + минимум ещё одна цена,
    # - либо две и более позиций с "голыми" ценами.
    if len(out) < 2:
        return None
    if explicit_count >= 1 and len(out) >= 2:
        return out
    if explicit_count == 0 and trailing_count >= 2:
        return out
    return None


def split_comma_list_with_single_price(line: str) -> list[tuple[str, float]] | None:
    """
    Строки вида:
      "Компот, кисель, каркадэ — 30.00 ₽"
      "Кетчуп, масло, сметана ... — 10.00 ₽"

    Одна цена в конце применяется к каждому элементу списка.
    """
    single = _parse_line(line)
    if single is None:
        return None
    # Не режем строки формата "... 0,5" без явных маркеров цены.
    if not re.search(r"(?:[—–-]\s*\d)|₽|руб|р\.?\b", line, re.IGNORECASE):
        return None
    name, price = single
    if "," not in name and ";" not in name:
        return None

    name_masked, dc_toks = _protect_volume_decimal_commas(name)
    raw_parts = _split_top_level_by_commas(name_masked)
    if len(raw_parts) < 2:
        return None

    cleaned_parts: list[str] = []
    for part in raw_parts:
        p = _restore_volume_decimal_commas(part.strip(), dc_toks)
        if not p:
            continue
        # Убираем шум после вычитки из docx-таблиц (висячие тире/запятые).
        p = p.strip(" \t\r\n,;:—–-")
        p = re.sub(r"\s+", " ", p).strip()
        if not p:
            continue
        # Слишком длинные куски чаще всего являются описанием одного блюда, а не списком.
        if len(p.split()) > 5:
            return None
        cleaned_parts.append(p)

    if len(cleaned_parts) < 2:
        return None

    out: list[tuple[str, float]] = []
    part_count = len(cleaned_parts)
    for i, p in enumerate(cleaned_parts):
        parsed = _parse_line(p)
        if parsed is None:
            out.append((p, price))
            continue
        own_name, own_price = parsed

        # Если цена в части явно размечена (рубли / "50-00"), считаем её собственной.
        if _has_explicit_price_marker(p):
            out.append((own_name, own_price))
            continue

        # Голое число в конце части (например "гренки 20") часто означает граммовку.
        base_name, trailing_num = _split_name_and_trailing_number_token(p)
        if trailing_num is not None:
            # Для коротких списков чаще всего это именно отдельная цена позиции:
            # "Хот-дог 80, Сочень — 50.00 ₽".
            if part_count <= 3:
                out.append((own_name, own_price))
                continue
            # Для длинных списков при общей цене в конце считаем это весом/граммовкой.
            out.append((base_name, price))
            continue

        out.append((p, price))
    return out if len(out) >= 2 else None


def _parse_one_line_to_items(line: str) -> tuple[list[tuple[str, float]], bool]:
    multi = split_multi_price_line(line)
    if multi is not None and len(multi) >= 2:
        return multi, True

    multi_comma_mixed = split_multi_comma_mixed_price_line(line)
    if multi_comma_mixed is not None and len(multi_comma_mixed) >= 2:
        return multi_comma_mixed, True

    multi_comma = split_multi_comma_price_line(line)
    if multi_comma is not None and len(multi_comma) >= 2:
        return multi_comma, True

    comma_single_price = split_comma_list_with_single_price(line)
    if comma_single_price is not None and len(comma_single_price) >= 2:
        return comma_single_price, True

    single = _parse_line(line)
    if single:
        return [single], False
    if multi is not None and len(multi) == 1:
        return multi, False
    return [], False


def _parse_line(line: str) -> tuple[str, float] | None:
    s = line.strip()
    if not s or len(s) < 2:
        return None
    s = re.sub(r"\s*[—–]\s*", " ", s)
    m = _PRICE_RE.match(s)
    if m:
        name = m.group("name").strip().rstrip(".,;")
        price = _parse_price_token(m.group("price"))
        if name and price is not None and not re.search(r"\d-$", name):
            return name, price

    parts = s.rsplit(None, 1)
    if len(parts) == 2:
        name = parts[0].strip()
        price = _parse_price_token(parts[1])
        if name and price is not None:
            return name, price
    return None


def parse_docx_bytes(data: bytes) -> list[tuple[str, float, DishKind]]:
    from io import BytesIO

    doc = Document(BytesIO(data))
    return _parse_docx_document(doc)


def _parse_docx_document(doc: Document) -> list[tuple[str, float, DishKind]]:
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if not cells:
                continue

            parsed_cells: list[tuple[str, float]] = []
            for cell in cells:
                if not re.search(r"[A-Za-zА-Яа-яЁё]", cell):
                    continue
                parsed = _parse_line(cell)
                if parsed is None:
                    continue
                name, price = parsed
                if name and price is not None:
                    parsed_cells.append((name, price))

            if len(parsed_cells) >= 2:
                for name, price in parsed_cells:
                    lines.append(f"{name} {price}")
                continue

            price = _parse_price_token(cells[-1] if cells else "")
            if len(cells) >= 2 and price is not None:
                name = " ".join(cells[:-1])
                if name:
                    lines.append(f"{name} {price}")
                    continue
            lines.append(" ".join(cells))

    out: list[tuple[str, float, DishKind]] = []
    seen: set[tuple[str, str]] = set()
    for line in lines:
        raw_items, multi_dash = _parse_one_line_to_items(line)
        for name, price in raw_items:
            name = sanitize_dish_name(name)
            name = _maybe_strip_category_prefix(name, multi_item_line=multi_dash)
            name = sanitize_dish_name(name)
            if not name:
                continue
            key = (name.casefold(), f"{price:.4f}")
            if key in seen:
                continue
            seen.add(key)
            kind = classify_dish(name)
            out.append((name, price, kind))
    return out


def parse_docx_path(path: Path) -> list[tuple[str, float, DishKind]]:
    return _parse_docx_document(Document(str(path)))
