from io import BytesIO

from docx import Document

from src.menu_parse import (
    _maybe_strip_category_prefix,
    _parse_line,
    _parse_one_line_to_items,
    classify_dish,
    parse_docx_bytes,
    sanitize_dish_name,
    strip_calories_from_dish_name,
    _parse_docx_document,
)


def test_parse_line_simple() -> None:
    assert _parse_line("Borsch 120") == ("Borsch", 120.0)
    assert _parse_line("Cutlet Kiev-style 150.50") == ("Cutlet Kiev-style", 150.5)
    assert _parse_line("Борщ 50-00") == ("Борщ", 50.0)
    assert _parse_line("Плов 100-5") == ("Плов", 100.5)


def test_classify_garnish_main() -> None:
    assert classify_dish("Гречка с маслом") in ("garnish", "main", "other")
    assert classify_dish("Котлета куриная") == "main"


def test_strip_calories() -> None:
    assert strip_calories_from_dish_name("Борщ 250 ккал") == "Борщ"
    assert strip_calories_from_dish_name("Салат (180 kcal)") == "Салат"
    assert strip_calories_from_dish_name("Плов 89 ккал") == "Плов"


def test_parse_line_strips_calories_via_export_path() -> None:
    # как в строке после склейки ячеек таблицы «название + ккал + цена»
    assert _parse_line("Котлета 220 ккал 150") == ("Котлета 220 ккал", 150.0)
    assert strip_calories_from_dish_name("Котлета 220 ккал") == "Котлета"


def test_sanitize_removes_nutrition_tail() -> None:
    raw = (
        "Салат «Летний» ( огурец, перец, помидор, лук, зелень,редис,майонез ;масло) "
        "100 198 6 16,7 7"
    )
    assert sanitize_dish_name(raw) == (
        "Салат «Летний» ( огурец, перец, помидор, лук, зелень,редис,майонез ;масло)"
    )


def test_sanitize_keeps_single_trailing_volume() -> None:
    # одно число в конце (0,5 л) — не хвост КБЖУ
    assert sanitize_dish_name("Напитки Компот,кисель 0,5") == "Напитки Компот,кисель 0,5"


def test_one_line_several_dash_prices() -> None:
    line = (
        "Выпечка хлеб 4-00 хлеб бор. 5-00 бутерброды 40-00 зелень 10-00 лимон 10-00"
    )
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    assert len(items) == 5
    cleaned = []
    for name, price in items:
        n = sanitize_dish_name(name)
        n = _maybe_strip_category_prefix(n, multi_item_line=True)
        n = sanitize_dish_name(n)
        cleaned.append((n, price))
    assert cleaned[0][0] == "хлеб"
    assert cleaned[0][1] == 4.0
    assert "бор" in cleaned[1][0].casefold()
    assert cleaned[2][0].casefold().startswith("бутерброд")


def test_one_line_several_comma_rub_prices() -> None:
    line = "Пицца 100р, Чебурек 70р, Сосиска в тесте 60р"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    assert len(items) == 3

    cleaned = []
    for name, price in items:
        n = sanitize_dish_name(name)
        n = _maybe_strip_category_prefix(n, multi_item_line=True)
        cleaned.append((n, price))

    assert cleaned[0] == ("Пицца", 100.0)
    assert "Чебурек" in cleaned[1][0]
    assert cleaned[2][1] == 60.0


def test_one_line_number_prefix_is_ignored() -> None:
    line = "4. Пицца 100р, Чебурек 70р"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    assert len(items) == 2
    names = [sanitize_dish_name(n) for n, _ in items]
    assert "Пицца" in names[0]
    assert "Чебурек" in names[1]


def test_one_line_comma_inside_parentheses_is_not_split() -> None:
    line = "Салат (Оливье, сытный) 100р, Борщ 80р"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    assert len(items) == 2

    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert "Оливье, сытный" in cleaned[0][0]
    assert cleaned[0][1] == 100.0
    assert cleaned[1][1] == 80.0


def test_one_line_category_prefix_is_stripped_for_first_item() -> None:
    line = "Выпечка Пицца 100р, Чебурек 70р"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    assert len(items) == 2

    cleaned = []
    for name, price in items:
        n = sanitize_dish_name(name)
        n = _maybe_strip_category_prefix(n, multi_item_line=True)
        cleaned.append((n, price))

    assert cleaned[0][0] == "Пицца"
    assert cleaned[1][0] == "Чебурек"


def test_one_line_does_not_split_without_explicit_rub_token() -> None:
    # "0,5" — это объём, а не цена; без `р/руб` не должно сработать разбиение.
    line = "Напитки Компот,кисель 0,5"
    items, multi = _parse_one_line_to_items(line)
    assert multi is False
    assert len(items) <= 1


def test_garnish_category_prefix_stripped_for_single_price_line() -> None:
    line = "Гарниры   Картошка жареная   60-00"
    items, multi = _parse_one_line_to_items(line)
    assert multi is False
    assert len(items) == 1
    name, price = items[0]
    cleaned = _maybe_strip_category_prefix(sanitize_dish_name(name), multi_item_line=multi)
    assert cleaned == "Картошка жареная"
    assert price == 60.0


def test_price_token_with_trailing_comma_is_parsed() -> None:
    # Пунктуация может "прилипнуть" к последней токенизированной ячейке из DOCX.
    line = "Гарниры Картошка жареная 60-00 ,"
    items, multi = _parse_one_line_to_items(line)
    assert len(items) == 1
    name, price = items[0]
    cleaned = _maybe_strip_category_prefix(sanitize_dish_name(name), multi_item_line=multi)
    assert cleaned == "Картошка жареная"
    assert price == 60.0


def test_parse_docx_table_comma_separated_items() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "4. Пицца 100р, Чебурек 70р, Сосиска в тесте 60р"

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    items = parse_docx_bytes(data)
    pairs = {(name, price) for name, price, _ in items}

    assert ("Пицца", 100.0) in pairs
    assert ("Чебурек", 70.0) in pairs
    assert ("Сосиска в тесте", 60.0) in pairs


def test_parse_docx_table_multiple_item_cells() -> None:
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Пицца 100р"
    table.cell(0, 1).text = "Чебурек 70р"

    buf = BytesIO()
    doc.save(buf)
    data = buf.getvalue()

    items = parse_docx_bytes(data)
    pairs = {(name, price) for name, price, _ in items}

    assert ("Пицца", 100.0) in pairs
    assert ("Чебурек", 70.0) in pairs


def test_comma_list_with_single_trailing_price_beverages() -> None:
    line = "Компот , кисель, каркадэ — 30.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Компот", 30.0) in cleaned
    assert ("кисель", 30.0) in cleaned
    assert ("каркадэ", 30.0) in cleaned


def test_comma_list_with_single_trailing_price_condiments() -> None:
    line = "Кетчуп, масло, сметана, майонез, , сгущенное молоко, гренки 20 — 10.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Кетчуп", 10.0) in cleaned
    assert ("масло", 10.0) in cleaned
    assert ("сметана", 10.0) in cleaned
    assert ("майонез", 10.0) in cleaned
    assert ("сгущенное молоко", 10.0) in cleaned
    assert ("гренки", 10.0) in cleaned


def test_comma_list_mixed_item_and_shared_price() -> None:
    line = "Хот-дог 80, Сочень — 50.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Хот-дог", 80.0) in cleaned
    assert ("Сочень", 50.0) in cleaned


def test_comma_list_two_items_bare_prices() -> None:
    line = "Хот-дог 80, Сочень 50"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Хот-дог", 80.0) in cleaned
    assert ("Сочень", 50.0) in cleaned


def test_comma_list_mixed_explicit_and_bare_prices() -> None:
    line = (
        "Пицца 100р, Чебурек 70р, Сосиска в тесте 60, "
        "Плюшка с маком 20, Сочник творожный 90,"
    )
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Пицца", 100.0) in cleaned
    assert ("Чебурек", 70.0) in cleaned
    assert ("Сосиска в тесте", 60.0) in cleaned
    assert ("Плюшка с маком", 20.0) in cleaned
    assert ("Сочник творожный", 90.0) in cleaned


def test_comma_list_with_hyphen_price_and_volume() -> None:
    line = "Компот,кисель 0.5 - 90.00"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Компот", 90.0) in cleaned
    assert ("кисель 0.5", 90.0) in cleaned


def test_comma_list_volume_with_decimal_comma_not_split() -> None:
    """Запятая в 0,5 л — не разделитель списка (иначе получаются «кисель 0» и «5»)."""
    line = "Компот, кисель 0,5 — 90.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Компот", 90.0) in cleaned
    assert ("кисель 0,5", 90.0) in cleaned
    assert all(p == 90.0 for _, p in cleaned)


def test_single_price_comma_list_strips_trailing_weight_in_long_list() -> None:
    line = "Кетчуп, масло, сметана, майонез, , сгущенное молоко, гренки 20 — 10.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("гренки", 10.0) in cleaned
    assert ("гренки 20", 10.0) not in cleaned


def test_single_price_comma_list_splits_beverages() -> None:
    line = "Компот , кисель, каркадэ — 30.00 ₽"
    items, multi = _parse_one_line_to_items(line)
    assert multi is True
    cleaned = [(sanitize_dish_name(n), p) for n, p in items]
    assert ("Компот", 30.0) in cleaned
    assert ("кисель", 30.0) in cleaned
    assert ("каркадэ", 30.0) in cleaned


def test_parse_docx_recovers_price_from_split_nutrition_row() -> None:
    doc = Document()
    # Эмуляция оборванной строки из DOCX, когда КБЖУ попадает в отдельную строку.
    doc.add_paragraph("Щи со сметаной 250 65,9 2")
    doc.add_paragraph("1,9 9,3 70")

    items = _parse_docx_document(doc)
    pairs = {(name, price) for name, price, _ in items}
    assert ("Щи со сметаной", 70.0) in pairs
    assert all(not name.startswith("1,9 9,3") for name, _price in pairs)


def test_parse_docx_strips_leading_category_phrase() -> None:
    doc = Document()
    doc.add_paragraph("Первые блюда Окрошка на квасе со сметаной 120")
    items = _parse_docx_document(doc)
    assert ("Окрошка на квасе со сметаной", 120.0, "other") in items
